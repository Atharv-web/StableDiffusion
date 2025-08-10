from __future__ import annotations

from typing import List, Optional, Type

from pydantic import BaseModel, Field
from crewai.tools import BaseTool

from pathlib import Path

try:
    from docx import Document
except Exception as exc:  # pragma: no cover
    raise ImportError(
        "python-docx is required for DocxAnnotatorTool. Please install it via 'pip install python-docx'."
    ) from exc


class DocxAnnotatorInput(BaseModel):
    """Input schema for the DocxAnnotatorTool.

    - documents: list of absolute or relative paths to .docx files to annotate
    - issues: list of dict-like objects containing at least a 'snippet' and 'type'.
      Optionally may include 'reference' and 'note' for extra guidance.
    - output_dir: optional directory to place annotated files; created if missing
    """

    documents: List[str] = Field(..., description="Paths to .docx files to annotate.")
    issues: List[dict] = Field(
        default_factory=list,
        description=(
            "Verification issues. Each should contain keys like 'snippet', 'type', and optional 'reference'/'note'."
        ),
    )
    output_dir: Optional[str] = Field(
        default=None, description="Directory to write annotated .docx files into."
    )


class DocxAnnotatorTool(BaseTool):
    name: str = "Docx Annotator"
    description: str = (
        "Annotates .docx files by inserting visible inline comment paragraphs near matched snippets. "
        "If a snippet is not found, appends an annotation at the end under an 'Annotations' section."
    )

    args_schema: Type[BaseModel]= DocxAnnotatorInput

    def _run(self, documents: List[str], issues: List[dict], output_dir: Optional[str] = None) -> str:
        output_paths: List[str] = []

        resolved_output_dir = self._resolve_output_dir(output_dir)
        resolved_output_dir.mkdir(parents=True, exist_ok=True)

        for document_path_str in documents:
            document_path = Path(document_path_str).resolve()
            if not document_path.exists() or document_path.suffix.lower() != ".docx":
                continue

            annotated_document = Document(str(document_path))

            # Track whether we found any matching paragraph for placement
            for issue in issues:
                snippet = str(issue.get("snippet", "")).strip()
                issue_type = str(issue.get("type", "issue")).strip()
                reference = str(issue.get("reference", "")).strip()
                note = str(issue.get("note", "")).strip()

                annotation_text = self._format_annotation(issue_type, snippet, reference, note)

                placed = False
                if snippet:
                    for paragraph in annotated_document.paragraphs:
                        if snippet in paragraph.text:
                            self._insert_paragraph_after(paragraph, annotation_text)
                            placed = True
                            break

                if not placed:
                    # Append to the end under an 'Annotations' section
                    self._ensure_annotations_section(annotated_document)
                    annotated_document.add_paragraph(annotation_text)

            output_file = resolved_output_dir / f"{document_path.stem}_annotated{document_path.suffix}"
            annotated_document.save(str(output_file))
            output_paths.append(str(output_file))

        return ", ".join(output_paths) if output_paths else "No annotations written."

    @staticmethod
    def _resolve_output_dir(output_dir: Optional[str]) -> Path:
        if output_dir:
            return Path(output_dir).resolve()
        # Default to an 'outputs' directory at the project root (two levels up from this file)
        return Path(__file__).resolve().parents[3] / "outputs"

    @staticmethod
    def _format_annotation(issue_type: str, snippet: str, reference: str, note: str) -> str:
        parts: List[str] = ["[ADGM Verification]"]
        if issue_type:
            parts.append(f"Type: {issue_type}")
        if snippet:
            # Truncate very long snippets for readability
            truncated = snippet if len(snippet) <= 300 else f"{snippet[:300]}..."
            parts.append(f"Snippet: {truncated}")
        if reference:
            parts.append(f"Reference: {reference}")
        if note:
            parts.append(f"Note: {note}")
        return " | ".join(parts)

    @staticmethod
    def _insert_paragraph_after(paragraph, text: str) -> None:
        # Insert a new paragraph element right after the given paragraph using the underlying XML tree
        new_paragraph = paragraph._p.addnext(paragraph._p.__class__())
        # Convert the new element into a python-docx Paragraph by adding a run to avoid empty paragraph issues
        # Workaround: add to the document at the end, then move not directly supported; simpler approach:
        # Instead, append to the end via add_paragraph and accept that location is approximate if XML manipulation fails.
        try:
            # Try creating a new paragraph after using the parent body element
            para = paragraph._element
            body = para.getparent()
            new_para = para.__class__()
            body.insert(body.index(para) + 1, new_para)
            # Add text to the new paragraph by wrapping with a Document interface
            # Locate the document object
            doc = paragraph._parent
            idx = list(doc.paragraphs).index(paragraph) + 1
            # Note: python-docx doesn't expose insert; after insertion, accessing by index should include new paragraph
            doc.paragraphs[idx].add_run(text)
        except Exception:
            # Fallback: append at end
            paragraph._parent.add_paragraph(text)

    @staticmethod
    def _ensure_annotations_section(doc: Document) -> None:
        # Add a heading once if not present
        for p in doc.paragraphs:
            if p.text.strip().lower() == "annotations":
                return
        doc.add_paragraph("Annotations")


