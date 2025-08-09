import streamlit as st
import docx
import os
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from pypdf import PdfReader
import pandas as pd

st.title("Corporate Agent")

CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Business Plan",
        "Register of Members and Directors",
        "Declaration of Compliance",
    ]
}

def get_data_text(data_dir):
    text = ""
    for filename in os.listdir(data_dir):
        filepath = os.path.join(data_dir, filename)
        if filename.endswith(".pdf"):
            pdf_reader = PdfReader(filepath)
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif filename.endswith(".docx"):
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text
        elif filename.endswith(".csv"):
            df = pd.read_csv(filepath)
            text += df.to_string()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

process = st.selectbox("Select the legal process:", list(CHECKLISTS.keys()))

google_api_key = st.text_input("Enter your Google API Key", type="password")

if "vector_store" not in st.session_state:
    st.session_state.vector_store = None

if st.button("Create Vector Store"):
    if google_api_key:
        with st.spinner("Creating vector store..."):
            raw_text = get_data_text("ADGMdata")
            text_chunks = get_text_chunks(raw_text)
            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=google_api_key)
            st.session_state.vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
            st.success("Vector store created successfully!")
    else:
        st.error("Please enter your Google API Key.")

uploaded_files = st.file_uploader("Upload .docx files", type="docx", accept_multiple_files=True)

def parse_docx(file):
    document = docx.Document(file)
    text = ""
    for para in document.paragraphs:
        text += para.text + "\n"
    return text

import json

def get_conversational_chain():
    prompt_template = """
    You are a legal assistant specializing in ADGM regulations. Your task is to review the provided paragraph from a document against the ADGM legal framework.
    Based on the context from our knowledge base and the full document text, identify any red flags, inconsistencies, or missing information in the paragraph.
    If you find an issue, please provide the output in a JSON format with the following keys: "issue", "severity", "suggestion".
    If there are no issues, return an empty JSON object.

    Context:\n {context}?\n
    Full Document:\n {full_document}\n
    Paragraph:\n {paragraph}\n

    Answer (in JSON format):
    """
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3, google_api_key=google_api_key)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "full_document", "paragraph"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

import io
import zipfile

def user_input(uploaded_files, process, missing_docs):
    if st.session_state.vector_store:
        report = {
            "process": process,
            "documents_uploaded": len(uploaded_files),
            "required_documents": len(CHECKLISTS[process]),
            "missing_document": missing_docs,
            "issues_found": []
        }

        reviewed_docs = {}

        for uploaded_file in uploaded_files:
            document = docx.Document(uploaded_file)
            full_document_text = "\n".join([para.text for para in document.paragraphs])
            chain = get_conversational_chain()

            for para in document.paragraphs:
                if para.text.strip() == "":
                    continue

                docs = st.session_state.vector_store.similarity_search(para.text)
                response = chain({"input_documents": docs, "full_document": full_document_text, "paragraph": para.text}, return_only_outputs=True)

                try:
                    analysis = json.loads(response["output_text"])
                    if analysis and analysis.get("issue"):
                        comment_text = f"Issue: {analysis['issue']}\nSeverity: {analysis['severity']}\nSuggestion: {analysis['suggestion']}"
                        para.add_comment(comment_text)
                        report["issues_found"].append({
                            "document": uploaded_file.name,
                            "section": para.text[:50] + "...",
                            "issue": analysis['issue'],
                            "severity": analysis['severity'],
                            "suggestion": analysis['suggestion']
                        })
                except json.JSONDecodeError:
                    pass

            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            reviewed_docs[uploaded_file.name] = file_stream

        st.subheader("Analysis Report")
        st.json(report)

        st.subheader("Downloads")
        st.download_button(
            label="Download JSON Report",
            data=json.dumps(report, indent=4),
            file_name="analysis_report.json",
            mime="application/json"
        )

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for doc_name, doc_stream in reviewed_docs.items():
                zip_file.writestr(f"reviewed_{doc_name}", doc_stream.getvalue())

        st.download_button(
            label="Download All Reviewed Documents (.zip)",
            data=zip_buffer.getvalue(),
            file_name="reviewed_documents.zip",
            mime="application/zip"
        )
    else:
        st.error("Vector store not created. Please create it first.")

if uploaded_files:
    uploaded_filenames = [f.name for f in uploaded_files]
    required_docs = CHECKLISTS[process]
    missing_docs = [doc for doc in required_docs if not any(doc.lower() in fname.lower() for fname in uploaded_filenames)]

    if missing_docs:
        st.warning(f"Missing Documents for {process}:")
        for doc in missing_docs:
            st.write(f"- {doc}")

    for uploaded_file in uploaded_files:
        st.subheader(f"Document: {uploaded_file.name}")
        document_text = parse_docx(uploaded_file)
        st.write(document_text)

    if st.button("Analyze Documents"):
        with st.spinner("Analyzing documents..."):
            user_input(uploaded_files, process, missing_docs)
